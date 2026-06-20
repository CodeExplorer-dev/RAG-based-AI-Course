"""DOCX 课件解析器
使用 python-docx 提取文本，按标题样式（Heading 1/2/3）识别章节结构，
返回带章节信息的页面列表，保留文档层级关系。
"""

from docx import Document


class DOCXParser:
    """DOCX 解析：按标题层级提取，每组 (section_title, text)"""
    HEADING_STYLES = {'Heading 1', 'Heading 2', 'Heading 3'}

    def extract(self, file_path: str) -> list[dict]:
        doc = Document(file_path)
        sections = []
        current_heading = ''
        current_paras = []

        def flush():
            nonlocal current_heading, current_paras
            if current_paras:
                body = '\n'.join(current_paras)
                title = current_heading or '（前言）'
                sections.append({
                    'page_number': None,  # DOCX 无固定页码，不生成虚假页码
                    'text': f'【{title}】\n{body}',
                    'heading': title,
                })
            current_paras = []

        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            sn = p.style.name if p.style else ''
            if sn in self.HEADING_STYLES:
                flush()
                current_heading = t
            else:
                current_paras.append(t)

        flush()

        if not sections:
            ps = [p.text for p in doc.paragraphs if p.text.strip()]
            if ps:
                return [{'page_number': None, 'text': '\n\n'.join(ps), 'heading': None}]
            return []
        return sections
